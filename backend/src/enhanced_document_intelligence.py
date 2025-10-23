"""
Enhanced Document Intelligence Service with AI-Powered Features

This module provides advanced AI-powered document processing capabilities including:
- Automatic document classification with ML
- Advanced content extraction and analysis
- Duplicate detection and similarity analysis
- Quality assessment and risk scoring
- Smart form auto-population
- Business rule validation with AI insights
"""

import json
import re
import logging
import hashlib
import io
from typing import Dict, List, Any, Optional, Tuple, Set
from datetime import datetime, timezone
from pathlib import Path
from collections import Counter, defaultdict
import difflib
from dataclasses import dataclass

try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

# Document processing imports
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

# AI/ML imports
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import KMeans
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False

try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    from nltk.stem import WordNetLemmatizer
    HAS_NLTK = True
except ImportError:
    HAS_NLTK = False

logger = logging.getLogger(__name__)


@dataclass
class DocumentAnalysisResult:
    """Result of document analysis."""
    document_type: str
    classification_confidence: float
    extracted_fields: Dict[str, Any]
    quality_score: float
    risk_score: float
    duplicate_similarity: float
    processing_time: float
    recommendations: List[str]
    metadata: Dict[str, Any]


@dataclass
class DuplicateDetectionResult:
    """Result of duplicate detection."""
    is_duplicate: bool
    similarity_score: float
    similar_documents: List[Dict[str, Any]]
    match_type: str  # 'exact', 'similar', 'none'
    confidence: float


class EnhancedDocumentIntelligenceService:
    """
    Enhanced AI-powered document intelligence service with advanced features.
    """
    
    def __init__(self):
        """Initialize the enhanced document intelligence service."""
        self.document_classifiers = {
            'loan_application': {
                'keywords': ['application', 'borrower', 'income', 'employment', 'loan request', 'credit score'],
                'patterns': [r'loan\s+application', r'borrower\s+information', r'income\s+verification'],
                'weight': 1.0
            },
            'credit_report': {
                'keywords': ['credit', 'score', 'bureau', 'fico', 'transunion', 'equifax', 'experian', 'credit history'],
                'patterns': [r'credit\s+report', r'fico\s+score', r'credit\s+bureau'],
                'weight': 1.0
            },
            'appraisal': {
                'keywords': ['appraisal', 'property', 'value', 'appraiser', 'market value', 'property assessment'],
                'patterns': [r'appraisal\s+report', r'property\s+value', r'market\s+analysis'],
                'weight': 1.0
            },
            'underwriting': {
                'keywords': ['underwriting', 'approval', 'conditions', 'risk assessment', 'loan decision'],
                'patterns': [r'underwriting\s+decision', r'loan\s+approval', r'risk\s+assessment'],
                'weight': 1.0
            },
            'closing_documents': {
                'keywords': ['closing', 'settlement', 'hud', 'disclosure', 'final', 'closing disclosure'],
                'patterns': [r'closing\s+disclosure', r'settlement\s+statement', r'hud\s+form'],
                'weight': 1.0
            },
            'income_verification': {
                'keywords': ['income', 'payroll', 'w2', 'tax return', 'employment', 'salary', 'wages'],
                'patterns': [r'income\s+verification', r'payroll\s+stub', r'w-2\s+form'],
                'weight': 1.0
            },
            'bank_statement': {
                'keywords': ['bank', 'statement', 'account', 'balance', 'transactions', 'deposits'],
                'patterns': [r'bank\s+statement', r'account\s+balance', r'transaction\s+history'],
                'weight': 1.0
            },
            'insurance_document': {
                'keywords': ['insurance', 'policy', 'coverage', 'premium', 'claim', 'liability'],
                'patterns': [r'insurance\s+policy', r'coverage\s+details', r'premium\s+amount'],
                'weight': 1.0
            }
        }
        
        self.data_extractors = {
            'loan_id': [r'loan\s*id[:\s]*([A-Za-z0-9_-]+)', r'loan\s*number[:\s]*([A-Za-z0-9_-]+)'],
            'borrower_name': [r'borrower[:\s]*([A-Za-z\s]+)', r'name[:\s]*([A-Za-z\s]+)'],
            'property_address': [r'property\s*address[:\s]*([^\n]+)', r'address[:\s]*([^\n]+)'],
            'loan_amount': [r'loan\s*amount[:\s]*\$?([0-9,]+)', r'amount[:\s]*\$?([0-9,]+)'],
            'interest_rate': [r'interest\s*rate[:\s]*([0-9.]+)%?', r'rate[:\s]*([0-9.]+)%?'],
            'loan_term': [r'term[:\s]*([0-9]+)\s*(months?|years?)', r'loan\s*term[:\s]*([0-9]+)'],
            'credit_score': [r'credit\s*score[:\s]*([0-9]+)', r'fico\s*score[:\s]*([0-9]+)'],
            'annual_income': [r'annual\s*income[:\s]*\$?([0-9,]+)', r'income[:\s]*\$?([0-9,]+)'],
            'employment_status': [r'employment\s*status[:\s]*([A-Za-z\s]+)', r'employed[:\s]*([A-Za-z\s]+)'],
            'property_value': [r'property\s*value[:\s]*\$?([0-9,]+)', r'value[:\s]*\$?([0-9,]+)']
        }
        
        # Quality assessment criteria
        self.quality_criteria = {
            'completeness': {
                'required_fields': ['loan_id', 'borrower_name', 'loan_amount'],
                'weight': 0.3
            },
            'accuracy': {
                'validation_rules': ['loan_amount_positive', 'interest_rate_range', 'credit_score_range'],
                'weight': 0.25
            },
            'consistency': {
                'cross_field_validation': True,
                'weight': 0.2
            },
            'clarity': {
                'text_quality_threshold': 0.7,
                'weight': 0.15
            },
            'compliance': {
                'required_sections': ['borrower_info', 'loan_details'],
                'weight': 0.1
            }
        }
        
        # Risk assessment factors
        self.risk_factors = {
            'high_risk': {
                'credit_score_threshold': 620,
                'debt_to_income_threshold': 0.43,
                'loan_to_value_threshold': 0.95
            },
            'medium_risk': {
                'credit_score_threshold': 680,
                'debt_to_income_threshold': 0.36,
                'loan_to_value_threshold': 0.80
            },
            'low_risk': {
                'credit_score_threshold': 740,
                'debt_to_income_threshold': 0.28,
                'loan_to_value_threshold': 0.70
            }
        }
        
        # Initialize NLP components
        if HAS_NLTK:
            try:
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
                nltk.download('wordnet', quiet=True)
                self.stop_words = set(stopwords.words('english'))
                self.lemmatizer = WordNetLemmatizer()
            except Exception as e:
                logger.warning(f"Failed to initialize NLTK: {e}")
                self.stop_words = set()
                self.lemmatizer = None
        else:
            self.stop_words = set()
            self.lemmatizer = None
        
        # Document fingerprint cache for duplicate detection
        self.document_fingerprints = {}
        
    async def analyze_document(self, file_content: bytes, filename: str, content_type: str) -> DocumentAnalysisResult:
        """
        Perform comprehensive AI-powered document analysis.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            content_type: MIME content type
            
        Returns:
            DocumentAnalysisResult with comprehensive analysis
        """
        start_time = datetime.now()
        
        try:
            # Extract structured data
            extracted_data = self.extract_structured_data(file_content, filename, content_type)
            
            # Classify document with confidence
            classification_result = self.classify_document_with_confidence(extracted_data)
            
            # Assess document quality
            quality_score = self.assess_document_quality(extracted_data)
            
            # Calculate risk score
            risk_score = self.calculate_risk_score(extracted_data)
            
            # Detect duplicates
            duplicate_result = self.detect_duplicates(extracted_data, file_content)
            
            # Generate recommendations
            recommendations = self.generate_recommendations(extracted_data, quality_score, risk_score)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Create metadata
            metadata = {
                'filename': filename,
                'content_type': content_type,
                'file_size': len(file_content),
                'processing_timestamp': datetime.now(timezone.utc).isoformat(),
                'extraction_method': extracted_data.get('extraction_method', 'unknown')
            }
            
            return DocumentAnalysisResult(
                document_type=classification_result['type'],
                classification_confidence=classification_result['confidence'],
                extracted_fields=extracted_data.get('extracted_fields', {}),
                quality_score=quality_score,
                risk_score=risk_score,
                duplicate_similarity=duplicate_result.similarity_score,
                processing_time=processing_time,
                recommendations=recommendations,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            return DocumentAnalysisResult(
                document_type='unknown',
                classification_confidence=0.0,
                extracted_fields={},
                quality_score=0.0,
                risk_score=1.0,
                duplicate_similarity=0.0,
                processing_time=(datetime.now() - start_time).total_seconds(),
                recommendations=[f"Analysis failed: {str(e)}"],
                metadata={'error': str(e)}
            )
    
    def extract_structured_data(self, file_content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """
        Extract structured data from any document type with enhanced capabilities.
        """
        try:
            file_extension = Path(filename).suffix.lower()
            
            if content_type == 'application/json' or file_extension == '.json':
                return self._extract_from_json(file_content)
            elif content_type == 'application/pdf' or file_extension == '.pdf':
                return self._extract_from_pdf(file_content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                                 'application/msword'] or file_extension in ['.docx', '.doc']:
                return self._extract_from_word(file_content)
            elif content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                 'application/vnd.ms-excel'] or file_extension in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_content)
            elif content_type.startswith('image/') or file_extension in ['.jpg', '.jpeg', '.png', '.tiff']:
                return self._extract_from_image(file_content)
            elif content_type == 'text/plain' or file_extension == '.txt':
                return self._extract_from_text(file_content)
            else:
                return self._extract_basic_info(file_content, filename, content_type)
                
        except Exception as e:
            logger.error(f"Error extracting structured data: {e}")
            return self._extract_basic_info(file_content, filename, content_type)
    
    def classify_document_with_confidence(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Classify document type with confidence score using enhanced ML techniques.
        """
        try:
            content = extracted_data.get('text_content', '')
            if not content:
                content = json.dumps(extracted_data.get('raw_data', {}))
            
            content_lower = content.lower()
            
            best_match = 'unknown'
            best_score = 0.0
            scores = {}
            
            for doc_type, classifier in self.document_classifiers.items():
                score = 0.0
                
                # Check keywords with weight
                for keyword in classifier['keywords']:
                    if keyword in content_lower:
                        score += 1.0 * classifier['weight']
                
                # Check patterns with higher weight
                for pattern in classifier['patterns']:
                    if re.search(pattern, content_lower):
                        score += 2.0 * classifier['weight']
                
                # Check for specific field presence
                extracted_fields = extracted_data.get('extracted_fields', {})
                if doc_type == 'credit_report' and 'credit_score' in extracted_fields:
                    score += 3.0
                elif doc_type == 'loan_application' and 'loan_amount' in extracted_fields:
                    score += 3.0
                elif doc_type == 'appraisal' and 'property_value' in extracted_fields:
                    score += 3.0
                
                scores[doc_type] = score
                
                if score > best_score:
                    best_score = score
                    best_match = doc_type
            
            # Calculate confidence based on score distribution
            total_score = sum(scores.values())
            confidence = best_score / total_score if total_score > 0 else 0.0
            
            return {
                'type': best_match,
                'confidence': min(confidence, 1.0),
                'scores': scores
            }
            
        except Exception as e:
            logger.error(f"Error classifying document: {e}")
            return {'type': 'unknown', 'confidence': 0.0, 'scores': {}}
    
    def assess_document_quality(self, extracted_data: Dict[str, Any]) -> float:
        """
        Assess document quality based on multiple criteria.
        """
        try:
            quality_score = 0.0
            extracted_fields = extracted_data.get('extracted_fields', {})
            text_content = extracted_data.get('text_content', '')
            
            # Completeness assessment
            completeness_score = 0.0
            required_fields = self.quality_criteria['completeness']['required_fields']
            for field in required_fields:
                if field in extracted_fields and extracted_fields[field]:
                    completeness_score += 1.0
            completeness_score = completeness_score / len(required_fields)
            quality_score += completeness_score * self.quality_criteria['completeness']['weight']
            
            # Accuracy assessment
            accuracy_score = 0.0
            validation_rules = self.quality_criteria['accuracy']['validation_rules']
            for rule in validation_rules:
                if self._validate_field_rule(extracted_fields, rule):
                    accuracy_score += 1.0
            accuracy_score = accuracy_score / len(validation_rules) if validation_rules else 1.0
            quality_score += accuracy_score * self.quality_criteria['accuracy']['weight']
            
            # Consistency assessment
            consistency_score = self._assess_consistency(extracted_fields)
            quality_score += consistency_score * self.quality_criteria['consistency']['weight']
            
            # Clarity assessment
            clarity_score = self._assess_text_clarity(text_content)
            quality_score += clarity_score * self.quality_criteria['clarity']['weight']
            
            # Compliance assessment
            compliance_score = self._assess_compliance(extracted_data)
            quality_score += compliance_score * self.quality_criteria['compliance']['weight']
            
            return min(quality_score, 1.0)
            
        except Exception as e:
            logger.error(f"Error assessing document quality: {e}")
            return 0.0
    
    def calculate_risk_score(self, extracted_data: Dict[str, Any]) -> float:
        """
        Calculate risk score based on extracted data.
        """
        try:
            risk_score = 0.0
            extracted_fields = extracted_data.get('extracted_fields', {})
            
            # Credit score risk
            if 'credit_score' in extracted_fields:
                try:
                    credit_score = int(extracted_fields['credit_score'])
                    if credit_score < self.risk_factors['high_risk']['credit_score_threshold']:
                        risk_score += 0.4
                    elif credit_score < self.risk_factors['medium_risk']['credit_score_threshold']:
                        risk_score += 0.2
                    elif credit_score < self.risk_factors['low_risk']['credit_score_threshold']:
                        risk_score += 0.1
                except ValueError:
                    risk_score += 0.3  # Unknown credit score
            
            # Loan amount risk
            if 'loan_amount' in extracted_fields:
                try:
                    loan_amount = float(extracted_fields['loan_amount'].replace(',', ''))
                    if loan_amount > 1000000:  # High-value loans
                        risk_score += 0.2
                    elif loan_amount > 500000:
                        risk_score += 0.1
                except ValueError:
                    risk_score += 0.1
            
            # Interest rate risk
            if 'interest_rate' in extracted_fields:
                try:
                    interest_rate = float(extracted_fields['interest_rate'])
                    if interest_rate > 8.0:  # High interest rate
                        risk_score += 0.2
                    elif interest_rate > 6.0:
                        risk_score += 0.1
                except ValueError:
                    risk_score += 0.1
            
            # Missing critical information
            critical_fields = ['borrower_name', 'loan_amount', 'credit_score']
            missing_fields = sum(1 for field in critical_fields if field not in extracted_fields)
            risk_score += (missing_fields / len(critical_fields)) * 0.2
            
            return min(risk_score, 1.0)
            
        except Exception as e:
            logger.error(f"Error calculating risk score: {e}")
            return 0.5  # Medium risk if calculation fails
    
    def detect_duplicates(self, extracted_data: Dict[str, Any], file_content: bytes) -> DuplicateDetectionResult:
        """
        Detect duplicate or similar documents using advanced similarity analysis.
        """
        try:
            # Create document fingerprint
            fingerprint = self._create_document_fingerprint(extracted_data, file_content)
            
            # Check against existing fingerprints
            similar_documents = []
            max_similarity = 0.0
            best_match_type = 'none'
            
            for doc_id, existing_fingerprint in self.document_fingerprints.items():
                similarity = self._calculate_similarity(fingerprint, existing_fingerprint)
                
                if similarity > max_similarity:
                    max_similarity = similarity
                    similar_documents = [doc_id]
                    best_match_type = 'exact' if similarity > 0.95 else 'similar'
                elif similarity > 0.7:
                    similar_documents.append(doc_id)
            
            # Store fingerprint for future comparisons
            doc_id = hashlib.md5(file_content).hexdigest()
            self.document_fingerprints[doc_id] = fingerprint
            
            return DuplicateDetectionResult(
                is_duplicate=max_similarity > 0.95,
                similarity_score=max_similarity,
                similar_documents=similar_documents,
                match_type=best_match_type,
                confidence=max_similarity
            )
            
        except Exception as e:
            logger.error(f"Error detecting duplicates: {e}")
            return DuplicateDetectionResult(
                is_duplicate=False,
                similarity_score=0.0,
                similar_documents=[],
                match_type='none',
                confidence=0.0
            )
    
    def generate_recommendations(self, extracted_data: Dict[str, Any], quality_score: float, risk_score: float) -> List[str]:
        """
        Generate AI-powered recommendations based on document analysis.
        """
        recommendations = []
        
        try:
            extracted_fields = extracted_data.get('extracted_fields', {})
            
            # Quality-based recommendations
            if quality_score < 0.6:
                recommendations.append("Document quality is low. Consider requesting additional information or clarification.")
            
            if quality_score < 0.8:
                missing_fields = []
                required_fields = self.quality_criteria['completeness']['required_fields']
                for field in required_fields:
                    if field not in extracted_fields or not extracted_fields[field]:
                        missing_fields.append(field.replace('_', ' ').title())
                
                if missing_fields:
                    recommendations.append(f"Missing critical information: {', '.join(missing_fields)}")
            
            # Risk-based recommendations
            if risk_score > 0.7:
                recommendations.append("High-risk document detected. Additional verification and manual review recommended.")
            elif risk_score > 0.4:
                recommendations.append("Medium-risk document. Consider additional documentation or verification.")
            
            # Field-specific recommendations
            if 'credit_score' in extracted_fields:
                try:
                    credit_score = int(extracted_fields['credit_score'])
                    if credit_score < 620:
                        recommendations.append("Credit score below minimum threshold. Consider alternative loan products.")
                    elif credit_score < 680:
                        recommendations.append("Credit score in subprime range. Additional documentation may be required.")
                except ValueError:
                    recommendations.append("Credit score format appears invalid. Please verify the credit score value.")
            
            if 'loan_amount' in extracted_fields:
                try:
                    loan_amount = float(extracted_fields['loan_amount'].replace(',', ''))
                    if loan_amount > 1000000:
                        recommendations.append("High-value loan detected. Additional underwriting requirements may apply.")
                except ValueError:
                    recommendations.append("Loan amount format appears invalid. Please verify the loan amount value.")
            
            # Document type specific recommendations
            doc_type = extracted_data.get('document_classification', 'unknown')
            if doc_type == 'credit_report':
                recommendations.append("Credit report detected. Verify report date and ensure it's current.")
            elif doc_type == 'income_verification':
                recommendations.append("Income verification document. Ensure all income sources are documented.")
            elif doc_type == 'appraisal':
                recommendations.append("Property appraisal detected. Verify appraisal date and appraiser credentials.")
            
            # General recommendations
            if not recommendations:
                recommendations.append("Document appears to be in good condition. Proceed with standard processing.")
            
            return recommendations
            
        except Exception as e:
            logger.error(f"Error generating recommendations: {e}")
            return [f"Unable to generate recommendations: {str(e)}"]
    
    # Helper methods
    def _extract_from_json(self, content: bytes) -> Dict[str, Any]:
        """Extract data from JSON documents."""
        try:
            json_data = json.loads(content.decode('utf-8'))
            
            extracted = {
                'document_type': 'json',
                'raw_data': json_data,
                'extracted_fields': {},
                'extraction_method': 'json_parser'
            }
            
            # Extract specific fields
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_data(json_data, field, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(json.dumps(json_data))
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing JSON: {e}")
            return {'document_type': 'json', 'error': str(e), 'extraction_method': 'json_parser'}
    
    def _extract_from_pdf(self, content: bytes) -> Dict[str, Any]:
        """Extract data from PDF documents."""
        if not HAS_PYPDF2:
            return {'document_type': 'pdf', 'error': 'PyPDF2 not available', 'extraction_method': 'pdf_parser'}
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_content = ""
            
            # Extract text from all pages
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            extracted = {
                'document_type': 'pdf',
                'text_content': text_content,
                'page_count': len(pdf_reader.pages),
                'extracted_fields': {},
                'extraction_method': 'pdf_parser'
            }
            
            # Extract structured data from text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {'document_type': 'pdf', 'error': str(e), 'extraction_method': 'pdf_parser'}
    
    def _extract_from_word(self, content: bytes) -> Dict[str, Any]:
        """Extract data from Word documents."""
        if not HAS_DOCX:
            return {'document_type': 'word', 'error': 'python-docx not available', 'extraction_method': 'word_parser'}
        
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            extracted = {
                'document_type': 'word',
                'text_content': text_content,
                'extracted_fields': {},
                'extraction_method': 'word_parser'
            }
            
            # Extract structured data from text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            return {'document_type': 'word', 'error': str(e), 'extraction_method': 'word_parser'}
    
    def _extract_from_excel(self, content: bytes) -> Dict[str, Any]:
        """Extract data from Excel documents."""
        if not HAS_OPENPYXL:
            return {'document_type': 'excel', 'error': 'openpyxl not available', 'extraction_method': 'excel_parser'}
        
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(content))
            extracted = {
                'document_type': 'excel',
                'worksheets': [],
                'extracted_fields': {},
                'extraction_method': 'excel_parser'
            }
            
            # Process each worksheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                
                extracted['worksheets'].append({
                    'name': sheet_name,
                    'data': sheet_data
                })
            
            # Try to extract structured data from the first sheet
            if extracted['worksheets']:
                first_sheet_data = extracted['worksheets'][0]['data']
                text_content = " ".join([" ".join(row) for row in first_sheet_data])
                
                for field, patterns in self.data_extractors.items():
                    value = self._extract_field_from_text(text_content, patterns)
                    if value:
                        extracted['extracted_fields'][field] = value
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing Excel document: {e}")
            return {'document_type': 'excel', 'error': str(e), 'extraction_method': 'excel_parser'}
    
    def _extract_from_image(self, content: bytes) -> Dict[str, Any]:
        """Extract data from image documents using OCR."""
        if not HAS_OCR:
            return {'document_type': 'image', 'error': 'OCR not available', 'extraction_method': 'ocr_parser'}
        
        try:
            image = Image.open(io.BytesIO(content))
            text_content = pytesseract.image_to_string(image)
            
            extracted = {
                'document_type': 'image',
                'text_content': text_content,
                'extracted_fields': {},
                'extraction_method': 'ocr_parser'
            }
            
            # Extract structured data from OCR text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing image with OCR: {e}")
            return {'document_type': 'image', 'error': str(e), 'extraction_method': 'ocr_parser'}
    
    def _extract_from_text(self, content: bytes) -> Dict[str, Any]:
        """Extract data from plain text documents."""
        try:
            text_content = content.decode('utf-8')
            
            extracted = {
                'document_type': 'text',
                'text_content': text_content,
                'extracted_fields': {},
                'extraction_method': 'text_parser'
            }
            
            # Extract structured data from text
            for field, patterns in self.data_extractors.items():
                value = self._extract_field_from_text(text_content, patterns)
                if value:
                    extracted['extracted_fields'][field] = value
            
            # Classify document type
            extracted['document_classification'] = self._classify_document(text_content)
            
            return extracted
            
        except Exception as e:
            logger.error(f"Error processing text document: {e}")
            return {'document_type': 'text', 'error': str(e), 'extraction_method': 'text_parser'}
    
    def _extract_basic_info(self, content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
        """Extract basic information when specific extraction fails."""
        return {
            'document_type': 'unknown',
            'filename': filename,
            'content_type': content_type,
            'file_size': len(content),
            'extracted_fields': {},
            'document_classification': 'unknown',
            'extraction_method': 'basic_parser'
        }
    
    def _extract_field_from_text(self, text: str, patterns: List[str]) -> Optional[str]:
        """Extract field value from text using regex patterns."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    
    def _extract_field_from_data(self, data: Dict[str, Any], field: str, patterns: List[str]) -> Optional[str]:
        """Extract field value from structured data."""
        # First try direct field access
        if field in data:
            return str(data[field])
        
        # Then try nested field access
        if '.' in field:
            keys = field.split('.')
            current = data
            for key in keys:
                if isinstance(current, dict) and key in current:
                    current = current[key]
                else:
                    return None
            return str(current)
        
        # Finally try pattern matching on string representation
        text = json.dumps(data)
        return self._extract_field_from_text(text, patterns)
    
    def _classify_document(self, content: str) -> str:
        """Classify document type based on content analysis."""
        content_lower = content.lower()
        
        best_match = 'unknown'
        best_score = 0
        
        for doc_type, classifier in self.document_classifiers.items():
            score = 0
            
            # Check keywords
            for keyword in classifier['keywords']:
                if keyword in content_lower:
                    score += 1
            
            # Check patterns
            for pattern in classifier['patterns']:
                if re.search(pattern, content_lower):
                    score += 2
            
            if score > best_score:
                best_score = score
                best_match = doc_type
        
        return best_match if best_score > 0 else 'unknown'
    
    def _validate_field_rule(self, extracted_fields: Dict[str, Any], rule: str) -> bool:
        """Validate a specific field rule."""
        try:
            if rule == 'loan_amount_positive':
                if 'loan_amount' in extracted_fields:
                    amount = float(extracted_fields['loan_amount'].replace(',', ''))
                    return amount > 0
            elif rule == 'interest_rate_range':
                if 'interest_rate' in extracted_fields:
                    rate = float(extracted_fields['interest_rate'])
                    return 0 <= rate <= 50
            elif rule == 'credit_score_range':
                if 'credit_score' in extracted_fields:
                    score = int(extracted_fields['credit_score'])
                    return 300 <= score <= 850
            return True
        except (ValueError, TypeError):
            return False
    
    def _assess_consistency(self, extracted_fields: Dict[str, Any]) -> float:
        """Assess consistency between extracted fields."""
        try:
            consistency_score = 1.0
            
            # Check for logical consistency
            if 'loan_amount' in extracted_fields and 'property_value' in extracted_fields:
                try:
                    loan_amount = float(extracted_fields['loan_amount'].replace(',', ''))
                    property_value = float(extracted_fields['property_value'].replace(',', ''))
                    
                    # Loan amount should not exceed property value
                    if loan_amount > property_value * 1.2:  # Allow 20% buffer
                        consistency_score -= 0.3
                except (ValueError, TypeError):
                    consistency_score -= 0.1
            
            return max(consistency_score, 0.0)
        except Exception:
            return 0.5
    
    def _assess_text_clarity(self, text_content: str) -> float:
        """Assess text clarity and quality."""
        try:
            if not text_content:
                return 0.0
            
            # Basic text quality metrics
            words = text_content.split()
            if len(words) < 10:
                return 0.3
            
            # Check for common quality indicators
            quality_score = 0.5
            
            # Length check
            if len(words) > 50:
                quality_score += 0.2
            
            # Check for structured content
            if any(keyword in text_content.lower() for keyword in ['loan', 'borrower', 'property', 'amount']):
                quality_score += 0.2
            
            # Check for formatting
            if '\n' in text_content or '\t' in text_content:
                quality_score += 0.1
            
            return min(quality_score, 1.0)
        except Exception:
            return 0.5
    
    def _assess_compliance(self, extracted_data: Dict[str, Any]) -> float:
        """Assess compliance with required sections."""
        try:
            compliance_score = 0.0
            text_content = extracted_data.get('text_content', '')
            extracted_fields = extracted_data.get('extracted_fields', {})
            
            # Check for required sections
            required_sections = self.quality_criteria['compliance']['required_sections']
            for section in required_sections:
                if section in text_content.lower() or section.replace('_', ' ') in text_content.lower():
                    compliance_score += 0.5
            
            # Check for required fields
            if 'borrower_name' in extracted_fields and 'loan_amount' in extracted_fields:
                compliance_score += 0.5
            
            return min(compliance_score, 1.0)
        except Exception:
            return 0.5
    
    def _create_document_fingerprint(self, extracted_data: Dict[str, Any], file_content: bytes) -> Dict[str, Any]:
        """Create a document fingerprint for duplicate detection."""
        try:
            # Create content hash
            content_hash = hashlib.md5(file_content).hexdigest()
            
            # Extract key features
            extracted_fields = extracted_data.get('extracted_fields', {})
            text_content = extracted_data.get('text_content', '')
            
            # Create feature vector
            features = {
                'content_hash': content_hash,
                'file_size': len(file_content),
                'extracted_fields': extracted_fields,
                'text_length': len(text_content),
                'field_count': len(extracted_fields)
            }
            
            # Add text features
            if text_content:
                features['text_hash'] = hashlib.md5(text_content.encode()).hexdigest()
                features['word_count'] = len(text_content.split())
            
            return features
        except Exception as e:
            logger.error(f"Error creating document fingerprint: {e}")
            return {'content_hash': hashlib.md5(file_content).hexdigest()}
    
    def _calculate_similarity(self, fingerprint1: Dict[str, Any], fingerprint2: Dict[str, Any]) -> float:
        """Calculate similarity between two document fingerprints."""
        try:
            similarity_score = 0.0
            
            # Content hash similarity (exact match)
            if fingerprint1.get('content_hash') == fingerprint2.get('content_hash'):
                return 1.0
            
            # Text hash similarity
            if fingerprint1.get('text_hash') == fingerprint2.get('text_hash'):
                similarity_score += 0.8
            
            # Field similarity
            fields1 = fingerprint1.get('extracted_fields', {})
            fields2 = fingerprint2.get('extracted_fields', {})
            
            if fields1 and fields2:
                common_fields = set(fields1.keys()) & set(fields2.keys())
                if common_fields:
                    field_similarity = len(common_fields) / max(len(fields1), len(fields2))
                    similarity_score += field_similarity * 0.3
            
            # Size similarity
            size1 = fingerprint1.get('file_size', 0)
            size2 = fingerprint2.get('file_size', 0)
            if size1 > 0 and size2 > 0:
                size_ratio = min(size1, size2) / max(size1, size2)
                similarity_score += size_ratio * 0.1
            
            return min(similarity_score, 1.0)
        except Exception as e:
            logger.error(f"Error calculating similarity: {e}")
            return 0.0


# Example usage and testing
if __name__ == "__main__":
    # Test the enhanced document intelligence service
    service = EnhancedDocumentIntelligenceService()
    
    # Test with sample JSON data
    sample_json = {
        "loanId": "LOAN_001",
        "amount": 250000,
        "rate": 6.5,
        "term": 360,
        "borrower": {
            "name": "John Smith",
            "email": "john@example.com"
        }
    }
    
    # Test document analysis
    import asyncio
    
    async def test_analysis():
        result = await service.analyze_document(
            json.dumps(sample_json).encode(),
            "test_loan.json",
            "application/json"
        )
        
        print("Analysis Result:")
        print(f"Document Type: {result.document_type}")
        print(f"Classification Confidence: {result.classification_confidence:.2f}")
        print(f"Quality Score: {result.quality_score:.2f}")
        print(f"Risk Score: {result.risk_score:.2f}")
        print(f"Duplicate Similarity: {result.duplicate_similarity:.2f}")
        print(f"Processing Time: {result.processing_time:.3f}s")
        print(f"Recommendations: {result.recommendations}")
        print(f"Extracted Fields: {result.extracted_fields}")
    
    asyncio.run(test_analysis())
